import os
from docx import Document
from openpyxl import Workbook

# Folder containing DOCX files
INPUT_FOLDER = "project"

# Output folder
OUTPUT_FOLDER = "output"

# Create output folder
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def convert_docx_to_excel(docx_path, excel_path):
    document = Document(docx_path)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Content"

    row = 1

    # Write paragraphs
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            sheet.cell(row=row, column=1).value = text
            row += 1

    # Write tables
    for table in document.tables:
        row += 1
        for table_row in table.rows:
            col = 1
            for cell in table_row.cells:
                sheet.cell(row=row, column=col).value = cell.text.strip()
                col += 1
            row += 1

    workbook.save(excel_path)


# Loop through all DOCX files
for filename in os.listdir(INPUT_FOLDER):

    if filename.lower().endswith(".docx"):

        input_path = os.path.join(INPUT_FOLDER, filename)

        output_name = os.path.splitext(filename)[0] + ".xlsx"

        output_path = os.path.join(OUTPUT_FOLDER, output_name)

        print(f"Converting {filename}...")

        convert_docx_to_excel(input_path, output_path)

print("\nAll files converted successfully!")